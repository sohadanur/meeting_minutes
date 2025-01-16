from pyannote.audio import Pipeline
import torchaudio
import os
import speech_recognition as sr
from docx import Document
from transformers import pipeline as hf_pipeline

# Instantiate the PyAnnote pipeline
pipeline = Pipeline.from_pretrained(
    "pyannote/speaker-diarization-3.1",
    use_auth_token=""
)

# Hugging Face summarization and sentiment-analysis pipelines
summarizer = hf_pipeline("summarization")
sentiment_analyzer = hf_pipeline("sentiment-analysis")

# Paths to files
VIDEO_PATH = r"/home/sohada/bacbon/videoplayback.mp4"
AUDIO_PATH = r"/home/sohada/bacbon/videoplayback.wav"
TRANSCRIPTION_DOCX = r"/home/sohada/bacbon/transcription.docx"

# Function to extract audio from video
def extract_audio(video_path, audio_path):
    print("Step 1: Extracting audio from video...")
    try:
        if not os.path.exists(audio_path):
            waveform, sample_rate = torchaudio.load(video_path)
            torchaudio.save(audio_path, waveform, sample_rate)
            print(f"Audio extracted and saved to: {audio_path}")
        else:
            print(f"Audio file already exists: {audio_path}")
    except Exception as e:
        print(f"Error during audio extraction: {e}")
        raise e

# Function to transcribe audio
def transcribe_audio(audio_path):
    recognizer = sr.Recognizer()
    transcription = ""
    try:
        print("Step 2: Transcribing audio...")
        with sr.AudioFile(audio_path) as source:
            audio = recognizer.record(source)
            transcription = recognizer.recognize_google(audio)
            print("Transcription completed.")
    except Exception as e:
        print(f"Error during transcription: {e}")
    return transcription

# Function to summarize text
def summarize_text(text):
    try:
        print("Step 3: Summarizing transcription...")
        summary = summarizer(text, max_length=130, min_length=30, do_sample=False)
        return summary[0]['summary_text']
    except Exception as e:
        print(f"Error during summarization: {e}")
        return "Summarization failed."

# Function to analyze sentiment
def analyze_sentiment(text):
    try:
        print("Step 4: Performing sentiment analysis...")
        sentiment = sentiment_analyzer(text)
        return sentiment
    except Exception as e:
        print(f"Error during sentiment analysis: {e}")
        return []

# Function to save data to .docx file
def save_to_docx(transcription, summary, sentiment, file_path):
    try:
        print("Step 5: Saving data to .docx file...")
        doc = Document()
        doc.add_heading("Speaker Diarization Transcription", level=1)
        doc.add_paragraph(transcription)
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)
        doc.add_heading("Sentiment Analysis", level=1)
        for result in sentiment:
            doc.add_paragraph(f"Text: {result.get('text', 'N/A')}\n"
                              f"Sentiment: {result['label']}\n"
                              f"Confidence: {result['score']:.2f}")
        doc.save(file_path)
        print(f"Data saved to: {file_path}")
    except Exception as e:
        print(f"Error while saving to .docx: {e}")

# Main function to process speaker diarization
def process_speaker_diarization():
    try:
        # Step 1: Extract audio
        extract_audio(VIDEO_PATH, AUDIO_PATH)

        # Step 2: Run the diarization pipeline
        print("Step 6: Running speaker diarization pipeline...")
        diarization = pipeline(AUDIO_PATH)
        print("Speaker Diarization Results:")
        diarized_segments = []
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            segment = f"Start: {round(turn.start, 2)}s, End: {round(turn.end, 2)}s, Speaker: {speaker}"
            diarized_segments.append(segment)
            print(segment)

        # Step 3: Transcribe audio
        transcription = transcribe_audio(AUDIO_PATH)

        # Step 4: Summarize transcription
        summary = summarize_text(transcription)

        # Step 5: Perform sentiment analysis
        sentiment = analyze_sentiment(transcription)

        # Step 6: Save data to .docx
        save_to_docx(transcription, summary, sentiment, TRANSCRIPTION_DOCX)

        print("Process completed successfully!")

    except FileNotFoundError as fnf_error:
        print(f"FileNotFoundError: {fnf_error}")
    except Exception as e:
        print(f"An error occurred: {e}")

# Run the script
if __name__ == "__main__":
    process_speaker_diarization()
